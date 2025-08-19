import os
from typing import List, Optional, Dict, Any

from mcp.server.fastmcp import FastMCP

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

mcp = FastMCP("word-writer")

# ===== Formatting defaults =====
FONT_LATIN = "Times New Roman"
FONT_EAST_ASIAN = "宋体"

SIZE_TITLE = 20
SIZE_BODY = 12
SIZE_H1 = 16
SIZE_H2 = 14
SIZE_H3 = 12

LINE_SPACING_BODY = 1.5  # 1.5 lines
INDENT_FIRST_LINE_CM = 0.74  # ~2个汉字缩进（约0.74cm）

MARGIN_CM = 2.5  # All sides


# ===== Helpers =====
def _ensure_dirs(filepath: str) -> None:
    os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)


def _apply_run_fonts(run, size_pt: int, bold: bool = False, italic: bool = False):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIAN)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def _configure_page_and_styles(doc: DocumentType):
    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

    # Normal style base config (Latin only; East Asian we set per-run)
    normal = doc.styles["Normal"]
    if normal.font:
        normal.font.name = FONT_LATIN
        normal.font.size = Pt(SIZE_BODY)

    # Heading styles (set Latin; East Asian will be enforced per-run)
    for style_name, sz in [("Heading 1", SIZE_H1), ("Heading 2", SIZE_H2), ("Heading 3", SIZE_H3)]:
        st = doc.styles[style_name]
        if st.font:
            st.font.name = FONT_LATIN
            st.font.size = Pt(sz)
            st.font.bold = True
        # Ensure H1 centered (冗余保障：函数写入时也会设置)
        if style_name == "Heading 1":
            try:
                st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    # Optional: configure Title style baseline (not strictly required since we set per-paragraph)
    try:
        title_style = doc.styles["Title"]
        if title_style.font:
            title_style.font.name = FONT_LATIN
            title_style.font.size = Pt(SIZE_TITLE)
            title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _load_or_create_document(file_path: str, author: Optional[str] = None) -> DocumentType:
    _ensure_dirs(file_path)
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()
        _configure_page_and_styles(doc)
        if author:
            doc.core_properties.author = author
        doc.save(file_path)
    return doc


def _save_document(doc: DocumentType, file_path: str):
    doc.save(file_path)


def _add_title(doc: DocumentType, text: str):
    # Document title paragraph (centered, large)
    p = doc.add_paragraph(style="Title") if "Title" in [s.name for s in doc.styles] else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _apply_run_fonts(run, SIZE_TITLE, bold=True)
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(24)
    # Also write core property
    try:
        doc.core_properties.title = text
    except Exception:
        pass
    return p


def _add_heading(doc: DocumentType, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    # enforce East Asian font and size per-level with bold
    size = SIZE_H1 if level == 1 else SIZE_H2 if level == 2 else SIZE_H3
    _apply_run_fonts(run, size, bold=True)
    # spacing
    fmt = p.paragraph_format
    if level == 1:
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(12)
    elif level == 2:
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
    else:
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
    return p


def _add_body_paragraph(
    doc: DocumentType,
    text: str,
    align: str = "left",
    first_line_indent: bool = True,
    spacing: float = LINE_SPACING_BODY,
):
    p = doc.add_paragraph(style="Normal")
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(INDENT_FIRST_LINE_CM)

    # 1.5 行距
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(text)
    _apply_run_fonts(run, SIZE_BODY, bold=False)
    return p


def _add_table(
    doc: DocumentType,
    headers: List[str],
    rows: List[List[str]],
    column_widths_cm: Optional[List[float]] = None,
    align: str = "center",
):
    cols = len(headers)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }.get(align.lower(), WD_TABLE_ALIGNMENT.CENTER)

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        _apply_run_fonts(run, SIZE_BODY, bold=True)
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # data rows
    for r in rows:
        row_cells = table.add_row().cells
        for i in range(cols):
            val = r[i] if i < len(r) else ""
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _apply_run_fonts(run, SIZE_BODY, bold=False)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # column widths
    if column_widths_cm:
        for col_idx, w in enumerate(column_widths_cm):
            if w and col_idx < cols:
                for row in table.rows:
                    row.cells[col_idx].width = Cm(w)

    return table


def _add_image(
    doc: DocumentType,
    image_path: str,
    width_cm: Optional[float] = None,
    caption: Optional[str] = None,
    align: str = "center",
):
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Image not found: {image_path}")

    # Add image
    width = Cm(width_cm) if width_cm else None
    pic = doc.add_picture(image_path, width=width)

    # Align the image by wrapping in a paragraph
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align.lower(), WD_ALIGN_PARAGRAPH.CENTER)

    # Caption (as a centered paragraph, small spacing)
    if caption:
        p = doc.add_paragraph()
        p.alignment = last_paragraph.alignment
        run = p.add_run(caption)
        _apply_run_fonts(run, SIZE_BODY, bold=False, italic=True)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(9)

    return pic


# ===== MCP Tools =====
@mcp.tool()
def init_document(file_path: str, title: Optional[str] = None, author: Optional[str] = None) -> Dict[str, Any]:
    """
    初始化或打开一个 .docx 文档，并设置页面边距与基础样式。
    - file_path: 文档路径（若不存在则创建）
    - title: 可选。若提供，将作为封面（居中，大号字体）
    - author: 可选。写入文档属性作者
    """
    doc = _load_or_create_document(file_path, author=author)
    # Ensure formatting every time (in case of existing doc created elsewhere)
    _configure_page_and_styles(doc)

    if title:
        _add_title(doc, title)

    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path}


@mcp.tool()
def write_title(file_path: str, text: str) -> Dict[str, Any]:
    """
    写入文档标题（封面标题）：20pt，加粗，居中，较大段前/段后间距；同时写入文档元数据 title。
    - file_path: 文档路径
    - text: 标题文本
    """
    doc = _load_or_create_document(file_path)
    _configure_page_and_styles(doc)  # 保证样式一致
    _add_title(doc, text)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "title": text}


@mcp.tool()
def write_heading_level_1(file_path: str, text: str) -> Dict[str, Any]:
    """写入一级标题（16pt，粗体，居中；中英文分别用 宋体/Times New Roman）"""
    doc = _load_or_create_document(file_path)
    _add_heading(doc, text, level=1)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "heading": text, "level": 1}


@mcp.tool()
def write_heading_level_2(file_path: str, text: str) -> Dict[str, Any]:
    """写入二级标题（14pt，粗体，左对齐；中英文分别用 宋体/Times New Roman）"""
    doc = _load_or_create_document(file_path)
    _add_heading(doc, text, level=2)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "heading": text, "level": 2}


@mcp.tool()
def write_heading_level_3(file_path: str, text: str) -> Dict[str, Any]:
    """写入三级标题（12pt，粗体，左对齐；中英文分别用 宋体/Times New Roman）"""
    doc = _load_or_create_document(file_path)
    _add_heading(doc, text, level=3)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "heading": text, "level": 3}


@mcp.tool()
def write_paragraph(
    file_path: str,
    text: str,
    align: str = "left",
    first_line_indent: bool = True,
) -> Dict[str, Any]:
    """
    写入正文段落（12pt，1.5倍行距，默认首行缩进2字符）。
    - align: left/center/right/justify
    - first_line_indent: 是否首行缩进
    """
    doc = _load_or_create_document(file_path)
    _add_body_paragraph(doc, text, align=align, first_line_indent=first_line_indent)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "chars": len(text)}


@mcp.tool()
def write_table(
    file_path: str,
    headers: List[str],
    rows: List[List[str]],
    column_widths_cm: Optional[List[float]] = None,
    align: str = "center",
) -> Dict[str, Any]:
    """
    写入表格数据（Table Grid风格，表头加粗，单元格居中）。
    - headers: 表头字段列表
    - rows: 行数据（二维数组）
    - column_widths_cm: 可选，各列宽度（cm）
    - align: left/center/right
    """
    doc = _load_or_create_document(file_path)
    _add_table(doc, headers, rows, column_widths_cm=column_widths_cm, align=align)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "rows": len(rows), "cols": len(headers)}


@mcp.tool()
def write_image(
    file_path: str,
    image_path: str,
    width_cm: Optional[float] = None,
    caption: Optional[str] = None,
    align: str = "center",
) -> Dict[str, Any]:
    """
    写入图片。
    - image_path: 本地图片路径
    - width_cm: 可选，宽度（cm）。不填则按图片原始尺寸
    - caption: 可选，图片标题/说明
    - align: left/center/right
    """
    doc = _load_or_create_document(file_path)
    _add_image(doc, image_path, width_cm=width_cm, caption=caption, align=align)
    _save_document(doc, file_path)
    return {"status": "ok", "file_path": file_path, "image": image_path}


@mcp.tool()
def get_word_writing_prompt() -> str:
    """
    获取：当用户要求写入Word文档时，推荐给模型遵循的工具调用流程提示词。
    """
    return (
        "当你需要将内容写入Word（.docx）文档时，请严格按以下流程进行工具调用：\n"
        "1) 首先调用 init_document：\n"
        "   - 参数：file_path（输出文件路径），可选 title（封面标题），author（作者）。\n"
        "   - 作用：创建/打开文档，并设置页边距、基础样式。\n"
        "   - 若需要在后续补充/修改标题，可调用 write_title。\n"
        "2) 按文档结构依次写入内容：\n"
        "   - 文档标题：使用 write_title 写入封面标题（居中）。\n"
        "   - 章节标题：使用 write_heading_level_1/2/3 写入不同层级标题（H1 居中，H2/H3 左对齐）。\n"
        "   - 正文：使用 write_paragraph 写入段落（默认首行缩进、1.5倍行距）。\n"
        "   - 表格：使用 write_table，传入 headers 与 rows（二维数组），必要时设置 column_widths_cm。\n"
        "   - 图片：使用 write_image，提供 image_path，必要时设置 width_cm 与 caption。\n"
        "3) 每次调用工具都会自动保存文档；请保持对同一个 file_path 的引用，避免分散到多个文件。\n"
        "4) 写作规范（默认已在工具中实现）：\n"
        "   - 字体：中文宋体，英文 Times New Roman。\n"
        "   - 标题：Title=20pt加粗居中；H1=16pt加粗居中；H2=14pt加粗左对齐；H3=12pt加粗左对齐；合适的段前/段后。\n"
        "   - 正文：12pt，1.5倍行距，默认首行缩进2字符。\n"
        "   - 表格：Table Grid 风格，表头加粗，内容居中；必要时设定列宽。\n"
        "   - 图片：支持插入本地图片，必要时设置宽度与说明。\n"
        "   - 公式：支持插入复杂 LaTeX 公式或简单 EQ 字段公式。\n"
        "5) 在输出中清楚说明你已完成了哪些写入（例如：已写入Title、H1、3段正文、1个表格、1张图片、2个公式），并提供最终文档路径。\n"
        "6) 若用户追加修改或继续写作，延续使用相同 file_path 继续写入即可。\n"
    )


if __name__ == "__main__":
    # Run the MCP server over stdio (default for FastMCP)
    mcp.run()